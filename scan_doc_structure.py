# -*- coding: utf-8 -*-
"""
Scan the document to find:
1. All image placeholder paragraphs (📌 Hình x.x)
2. Cover page content
3. All heading structure for image placement map
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.oxml.ns import qn

doc = Document('final report_FINAL_v6.docx')

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print()

# 1. Find all inline images already inserted
print("=== IMAGES ALREADY IN DOCUMENT ===")
img_count = 0
for i, para in enumerate(doc.paragraphs):
    # Check for inline images in runs
    for run in para.runs:
        drawings = run._r.findall('.//' + qn('a:blip'), 
                                   namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        if drawings or run._r.find('.//' + qn('w:drawing')) is not None:
            img_count += 1
            print(f"  [IMG in para {i}]: {para.text[:80] or '(empty para with image)'}")

print(f"Total images found: {img_count}")
print()

# 2. Find placeholder paragraphs (📌 Hình)
print("=== PLACEHOLDER IMAGE PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if ('📌' in t or '📌 Hình' in t or 'Hình' in t) and ('Screenshot' in t or 'Diagram' in t or 'ERD' in t or 'ETL' in t or 'DFD' in t or 'Activity' in t or 'Use Case' in t or 'Sequence' in t or 'Component' in t or 'State' in t):
        print(f"  [para {i}]: {t[:120]}")
print()

# 3. Cover page content
print("=== FIRST 20 PARAGRAPHS (COVER PAGE) ===")
for i, para in enumerate(doc.paragraphs[:25]):
    t = para.text.strip()
    s = para.style.name
    if t:
        print(f"  [{i}] [{s}]: {t[:100]}")
