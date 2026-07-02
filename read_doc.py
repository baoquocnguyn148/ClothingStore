# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document('final report_FINAL_v3.docx')
print(f'Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}, Sections: {len(doc.sections)}')
print()

for para in doc.paragraphs:
    text = para.text.strip()
    style = para.style.name
    if not text:
        continue
    if 'Heading 1' in style:
        print(f'\n=== H1: {text}')
    elif 'Heading 2' in style:
        print(f'\n  -- H2: {text}')
    elif 'Heading 3' in style:
        print(f'\n    - H3: {text}')
    elif 'Caption' in style or 'caption' in style.lower():
        print(f'    [CAPTION]: {text}')
    else:
        # Print first 300 chars of body text
        snippet = text[:300].replace('\n', ' ')
        print(f'       {snippet}')
