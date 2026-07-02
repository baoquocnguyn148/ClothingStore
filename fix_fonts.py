# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

IN_FILE  = "final report_FINAL.docx"
OUT_FILE = "final report_FINAL_v2.docx"

doc = Document(IN_FILE)

# ─── 1. Set the document theme fonts globally ─────────────────────────────
# Patch the document's theme fonts to Times New Roman (body) + Arial (heading)
# This is done via the document settings XML

BODY_FONT    = "Times New Roman"
HEADING_FONT = "Times New Roman"

def set_style_font(style, font_name, size_pt=None, bold=None, color_rgb=None, italic=None):
    font = style.font
    font.name = font_name
    # Also set East Asian and Complex Script font (critical for Vietnamese)
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    
    # East Asian font
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'),     font_name)
    rFonts.set(qn('w:hAnsi'),     font_name)
    rFonts.set(qn('w:cs'),        font_name)
    rFonts.set(qn('w:eastAsia'),  font_name)
    
    if size_pt:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if color_rgb:
        font.color.rgb = RGBColor(*color_rgb)

# ─── 2. Apply fonts to all styles ─────────────────────────────────────────
print("Setting fonts on styles...")

# Normal body text
set_style_font(doc.styles['Normal'], BODY_FONT, size_pt=13, bold=False)

# Headings
h1_color = (0x1E, 0x3A, 0x5F)   # dark navy
h2_color = (0x1D, 0x4E, 0xD8)   # blue
h3_color = (0x37, 0x47, 0x51)   # dark gray

set_style_font(doc.styles['Heading 1'], HEADING_FONT, size_pt=16, bold=True, color_rgb=h1_color)
set_style_font(doc.styles['Heading 2'], HEADING_FONT, size_pt=14, bold=True, color_rgb=h2_color)
set_style_font(doc.styles['Heading 3'], HEADING_FONT, size_pt=13, bold=True, color_rgb=h3_color)
set_style_font(doc.styles['Heading 4'], HEADING_FONT, size_pt=12, bold=True, italic=True, color_rgb=h3_color)

print("  ✓ Heading 1: Times New Roman 16pt Bold Navy")
print("  ✓ Heading 2: Times New Roman 14pt Bold Blue")
print("  ✓ Heading 3: Times New Roman 13pt Bold DarkGray")
print("  ✓ Normal: Times New Roman 13pt")

# ─── 3. Fix TOC styles ────────────────────────────────────────────────────
print("\nSetting TOC styles...")
toc_style_names = ['TOC 1', 'TOC 2', 'TOC 3', 'TOC 4', 'TOC Heading']
for name in toc_style_names:
    try:
        s = doc.styles[name]
        set_style_font(s, BODY_FONT, size_pt=12)
        print(f"  ✓ Set font on style: {name}")
    except:
        pass

# ─── 4. Fix all runs font inline (override any run-level font) ─────────────
print("\nFixing run-level font overrides...")
fixed = 0
bad_fonts = ['Calibri', 'Calibri (Body)', 'Arial', 'Cambria', 'Corbel', 'Candara',
             'Century Gothic', 'Trebuchet MS', 'Verdana', 'Tahoma', 'Courier New']

for para in doc.paragraphs:
    for run in para.runs:
        if run.font.name in bad_fonts or run.font.name is None:
            # Fix run-level font via XML directly
            rPr = run._r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run._r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                # Remove bad font tags so style inherits correctly
                rPr.remove(rFonts)
                fixed += 1
print(f"  ✓ Cleared {fixed} run-level font overrides")

# ─── 5. Fix heading spacing (before/after) ────────────────────────────────
print("\nSetting heading paragraph spacing...")

def set_style_spacing(style, before_pt, after_pt, line_rule=None, line_val=None):
    pPr = style.element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        style.element.append(pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'),  str(int(after_pt  * 20)))
    if line_val:
        spacing.set(qn('w:line'),     str(int(line_val * 240)))
        spacing.set(qn('w:lineRule'), 'auto')

set_style_spacing(doc.styles['Heading 1'], before_pt=24, after_pt=12, line_val=1.15)
set_style_spacing(doc.styles['Heading 2'], before_pt=18, after_pt=6,  line_val=1.15)
set_style_spacing(doc.styles['Heading 3'], before_pt=12, after_pt=4,  line_val=1.15)
set_style_spacing(doc.styles['Normal'],    before_pt=0,  after_pt=8,  line_val=1.5)
print("  ✓ Spacing set for all heading levels and Normal")

# ─── 6. Replace the bad TOC field with a clean one ─────────────────────────
print("\nReplacing TOC field with clean version...")

# Find the CONTENTS heading + the TOC field paragraph just after it
contents_para_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'CONTENTS' and 'Heading' in p.style.name:
        contents_para_idx = i
        break

if contents_para_idx is not None:
    # Remove old TOC paragraphs (the fldSimple we inserted before)
    # Look ahead up to 5 paragraphs and remove the one with fldSimple
    paras_el = doc.paragraphs[contents_para_idx]._element.getparent()
    contents_el = doc.paragraphs[contents_para_idx]._element
    
    # Find the next sibling that has fldSimple
    sibling = contents_el.getnext()
    if sibling is not None:
        fld = sibling.find('.//' + qn('w:fldSimple'))
        if fld is not None:
            sibling.getparent().remove(sibling)
            print("  ✓ Removed old TOC field paragraph")

    # Now insert a clean, properly formatted TOC
    # Create a paragraph with TOC field
    new_p_el = OxmlElement('w:p')
    
    # Paragraph properties (TOC style)
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    new_p_el.append(pPr)
    
    # fldSimple with TOC instruction
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), ' TOC \\o "1-3" \\h \\z \\u ')
    
    # Run inside fldSimple with placeholder text
    run_el = OxmlElement('w:r')
    rPr_el = OxmlElement('w:rPr')
    
    # Set font on run
    rFonts_el = OxmlElement('w:rFonts')
    rFonts_el.set(qn('w:ascii'),    BODY_FONT)
    rFonts_el.set(qn('w:hAnsi'),    BODY_FONT)
    rFonts_el.set(qn('w:cs'),       BODY_FONT)
    rFonts_el.set(qn('w:eastAsia'), BODY_FONT)
    rPr_el.append(rFonts_el)
    
    sz_el = OxmlElement('w:sz')
    sz_el.set(qn('w:val'), '24')  # 12pt
    rPr_el.append(sz_el)
    
    run_el.append(rPr_el)
    t_el = OxmlElement('w:t')
    t_el.text = 'Right-click this line → Update Field (or press Ctrl+A then F9 in Word)'
    run_el.append(t_el)
    fldSimple.append(run_el)
    
    new_p_el.append(fldSimple)
    
    # Insert right after CONTENTS heading
    contents_el.addnext(new_p_el)
    print("  ✓ Clean TOC field inserted")

# ─── 7. Set margin (A4 standard thesis: top/bottom 25mm, left 35mm, right 20mm) ──
print("\nSetting page margins...")
from docx.shared import Mm
for section in doc.sections:
    section.top_margin    = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin   = Mm(35)   # binding side
    section.right_margin  = Mm(20)
print("  ✓ Margins: Top 25mm | Bottom 25mm | Left 35mm | Right 20mm")

# ─── 8. Save ──────────────────────────────────────────────────────────────
print(f"\nSaving to '{OUT_FILE}'...")
doc.save(OUT_FILE)
print(f"\n✅ DONE — '{OUT_FILE}' saved successfully!")
print("\n📌 NEXT STEPS in Microsoft Word:")
print("   1. Open 'final report_FINAL_v2.docx'")
print("   2. Click on the TOC placeholder text")
print("   3. Press Ctrl+A → F9 → 'Update entire table'")
print("   4. TOC will auto-fill with proper page numbers and fonts")
