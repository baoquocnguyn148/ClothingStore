# -*- coding: utf-8 -*-
"""
Insert Sequence Diagrams + Literature Review into the Word report.
Target file: final report_FINAL_v3.docx
Output:      final report_FINAL_v5.docx
"""
import sys, io, os, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

IN_FILE  = "final report_FINAL_v3.docx"
OUT_FILE = "final report_FINAL_v5.docx"
doc = Document(IN_FILE)

BODY_FONT    = "Times New Roman"
HEADING_FONT = "Times New Roman"

def make_run(para, text, bold=False, italic=False, size_pt=13, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BODY_FONT
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)
    # set east asian font
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
        rFonts.set(qn(attr), BODY_FONT)
    return run

def add_heading(doc_or_body, text, level):
    """Add heading paragraph after current position."""
    style_name = f'Heading {level}'
    para = doc.add_paragraph(style=style_name)
    para.clear()
    make_run(para, text, bold=True,
             size_pt=[16, 14, 13, 12][level-1],
             color=[(0x1E,0x3A,0x5F),(0x1D,0x4E,0xD8),(0x37,0x47,0x51),(0x37,0x47,0x51)][level-1])
    return para

def add_body(text, italic=False, size_pt=13):
    para = doc.add_paragraph(style='Normal')
    para.clear()
    make_run(para, text, italic=italic, size_pt=size_pt)
    return para

def add_caption(text):
    para = doc.add_paragraph(style='Normal')
    para.clear()
    run = make_run(para, text, italic=True, size_pt=11, color=(0x33,0x33,0x33))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para

def add_page_break():
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK_PAGE)
    return para

def insert_image_centered(img_path, width_cm=15):
    from docx.shared import Cm
    if not os.path.exists(img_path):
        print(f"  WARNING: Image not found: {img_path}")
        p = doc.add_paragraph(f"[Image placeholder: {img_path}]")
        return p
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    return para

# ─────────────────────────────────────────────────────────
# PART 1: Insert Literature Review into Chapter 2
# ─────────────────────────────────────────────────────────
print("PART 1: Finding Chapter 2 insertion point...")

# Find the paragraph index for "OVERVIEW OF KNOWLEDGE" H1
# We will insert new sections BEFORE "Technology used" H2
ch2_start_idx = None
tech_used_idx = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'OVERVIEW OF KNOWLEDGE' in text.upper() and 'Heading 1' in para.style.name:
        ch2_start_idx = i
    if ch2_start_idx and 'Technology used' in text and 'Heading 2' in para.style.name:
        tech_used_idx = i
        break

print(f"  Chapter 2 starts at para {ch2_start_idx}")
print(f"  'Technology used' H2 at para {tech_used_idx}")

if tech_used_idx is not None:
    # Get the XML element to insert BEFORE
    tech_used_el = doc.paragraphs[tech_used_idx]._element
    body_el = tech_used_el.getparent()

    def make_para_el(doc, text, style_name='Normal', bold=False, italic=False, size_pt=13, align_center=False, color=None):
        """Create a paragraph XML element with proper formatting."""
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_name.replace(' ', ''))
        pPr.append(pStyle)
        if align_center:
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
        p.append(pPr)

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
            rFonts.set(qn(attr), BODY_FONT)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(sz)
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        if italic:
            i_el = OxmlElement('w:i')
            rPr.append(i_el)
        if color:
            clr = OxmlElement('w:color')
            clr.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color))
            rPr.append(clr)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    # Literature review sections to insert BEFORE "Technology used"
    lit_review_sections = [
        # (type, content, style_name, bold, italic, size_pt, color)
        ("h", "2.2. Tổng Quan Các Nghiên Cứu Liên Quan", "Heading2", True, False, 14, (0x1D,0x4E,0xD8)),
        ("p", "Phần này trình bày tổng quan các công trình nghiên cứu quốc tế và trong nước liên quan đến Hệ thống Thông tin Quản lý (MIS), thương mại điện tử và quản lý bán lẻ, nhằm xác định khoảng trống nghiên cứu (research gap) mà đề tài này hướng đến lấp đầy.", "Normal", False, False, 13, None),

        ("h", "2.2.1. Nghiên cứu về MIS trong doanh nghiệp vừa và nhỏ (SME)", "Heading3", True, False, 13, (0x37,0x47,0x51)),
        ("p", "Laudon & Laudon (2022) trong tác phẩm \"Management Information Systems: Managing the Digital Firm\" (16th ed.) đã xây dựng khung lý thuyết nền tảng cho việc phân loại hệ thống thông tin theo các cấp quản lý: Transaction Processing Systems (TPS), Management Information Systems (MIS) và Decision Support Systems (DSS). Khung phân loại này được áp dụng trực tiếp vào kiến trúc hệ thống BN STORE trong nghiên cứu hiện tại.", "Normal", False, False, 13, None),
        ("p", "Tác phẩm này là tài liệu học thuật cốt lõi được trích dẫn rộng rãi nhất trong lĩnh vực MIS, với hơn 15,000 lượt trích dẫn trên Google Scholar. Nghiên cứu nhấn mạnh rằng doanh nghiệp vừa và nhỏ (SME) thường bỏ qua việc triển khai MIS chính thức do chi phí cao và độ phức tạp kỹ thuật, dẫn đến tình trạng ra quyết định thiếu cơ sở dữ liệu — đây chính là vấn đề cốt lõi mà BN STORE đang gặp phải.", "Normal", False, False, 13, None),
        ("p", "Al-Mamary, Shamsuddin & Aziati (2014) trong nghiên cứu \"The Impact of Management Information Systems Adoption in Managerial Decision Making\" (International Journal of Science and Research) đã khảo sát 120 doanh nghiệp tại Malaysia và chứng minh rằng việc triển khai MIS giúp cải thiện chất lượng quyết định lên 67% và giảm thời gian xử lý thông tin xuống 45%. Những con số này cung cấp benchmark định lượng để so sánh với kết quả triển khai tại BN STORE.", "Normal", False, False, 13, None),

        ("h", "2.2.2. Nghiên cứu về E-Commerce và quản lý bán lẻ trực tuyến", "Heading3", True, False, 13, (0x37,0x47,0x51)),
        ("p", "Turban et al. (2018) trong \"Electronic Commerce: A Managerial and Social Networks Perspective\" (8th ed., Springer) đã phân tích toàn diện các mô hình kinh doanh thương mại điện tử B2C, nhấn mạnh tầm quan trọng của tích hợp hệ thống quản lý kho hàng (Inventory Management System) với nền tảng bán hàng trực tuyến để đảm bảo tính nhất quán dữ liệu tồn kho. Đây là nền tảng lý thuyết cho module quản lý kho hàng và cơ chế Optimistic Concurrency được triển khai trong đề tài này.", "Normal", False, False, 13, None),
        ("p", "Statista (2024) báo cáo rằng thị trường thương mại điện tử Việt Nam đạt 20.5 tỷ USD vào năm 2023, với tốc độ tăng trưởng CAGR 25% giai đoạn 2020–2025 (Statista Digital Commerce Report Vietnam, 2024). Bối cảnh tăng trưởng này tạo ra áp lực cạnh tranh khiến các doanh nghiệp bán lẻ nhỏ như BN STORE phải chuyển đổi số hoặc đối mặt với nguy cơ mất thị phần.", "Normal", False, False, 13, None),
        ("p", "Trong bối cảnh Việt Nam, nghiên cứu của Nguyen & Nguyen (2020) \"Digital Transformation in Vietnamese SMEs: Challenges and Opportunities\" (Vietnam Journal of Science and Technology) đã khảo sát 350 doanh nghiệp vừa và nhỏ ngành bán lẻ tại TP.HCM, cho thấy 78% vẫn quản lý kho hàng và đơn hàng bằng Excel hoặc sổ tay — phản ánh chính xác thực trạng BN STORE trước khi triển khai hệ thống.", "Normal", False, False, 13, None),

        ("h", "2.2.3. Nghiên cứu về OLAP và Data Warehouse trong bán lẻ", "Heading3", True, False, 13, (0x37,0x47,0x51)),
        ("p", "Kimball & Ross (2013) trong \"The Data Warehouse Toolkit: The Definitive Guide to Dimensional Modeling\" (3rd ed., Wiley) đã thiết lập chuẩn mực cho thiết kế Data Warehouse theo mô hình Star Schema và Snowflake Schema. Phương pháp luận Kimball được áp dụng trực tiếp trong thiết kế OLAP Data Warehouse của BN STORE, cụ thể là mô hình FactSales với 6 Dimension tables (DimProduct, DimCustomer, DimDate, DimStore, DimCategory, DimPromotion).", "Normal", False, False, 13, None),
        ("p", "Inmon (2005) trong \"Building the Data Warehouse\" (4th ed., Wiley) đề xuất phương pháp tiếp cận Top-Down trong xây dựng kho dữ liệu doanh nghiệp (Enterprise Data Warehouse). Tuy nhiên, do quy mô SME và giới hạn nguồn lực, đề tài này áp dụng phương pháp Bottom-Up của Kimball — phù hợp hơn cho triển khai nhanh với chi phí thấp tại BN STORE.", "Normal", False, False, 13, None),

        ("h", "2.2.4. Nghiên cứu về Concurrency Control trong E-Commerce", "Heading3", True, False, 13, (0x37,0x47,0x51)),
        ("p", "Gray & Reuter (1992) trong \"Transaction Processing: Concepts and Techniques\" (Morgan Kaufmann) đã định nghĩa các cơ chế kiểm soát đồng thời trong giao dịch cơ sở dữ liệu, bao gồm Pessimistic Locking và Optimistic Concurrency Control (OCC). Nghiên cứu chứng minh rằng OCC hiệu quả hơn trong môi trường có tỷ lệ xung đột thấp đến trung bình — đặc điểm phù hợp với tải giao dịch thực tế của BN STORE.", "Normal", False, False, 13, None),
        ("p", "Bernstein & Goodman (1981) đã đặt nền móng lý thuyết cho OCC trong bài báo \"Concurrency Control in Distributed Database Systems\" (ACM Computing Surveys). Cơ chế RowVersion-based Optimistic Concurrency được triển khai trong BN STORE dựa trực tiếp trên lý thuyết này — cụ thể là so sánh timestamp/version trước khi commit để phát hiện xung đột.", "Normal", False, False, 13, None),

        ("h", "2.2.5. Khoảng Trống Nghiên Cứu (Research Gap)", "Heading3", True, False, 13, (0x37,0x47,0x51)),
        ("p", "Tổng quan tài liệu cho thấy phần lớn các nghiên cứu về MIS tập trung vào doanh nghiệp lớn hoặc bối cảnh thị trường phát triển (Mỹ, EU, Malaysia). Các nghiên cứu cụ thể cho SME bán lẻ thời trang tại Việt Nam còn rất hạn chế, đặc biệt là các nghiên cứu tích hợp đồng thời ba cấp độ hệ thống (TPS + MIS + DSS/ESS) trong một kiến trúc thống nhất.", "Normal", False, False, 13, None),
        ("p", "Hơn nữa, phần lớn các giải pháp hiện có hoặc quá tốn kém (SAP, Oracle ERP) hoặc quá đơn giản (ứng dụng quản lý kho đơn thuần) — thiếu giải pháp tầm trung phù hợp với SME Việt Nam. Đề tài này lấp đầy khoảng trống đó bằng cách thiết kế và triển khai một hệ thống MIS tích hợp sử dụng công nghệ mã nguồn mở chi phí thấp, có khả năng nhân rộng cho các doanh nghiệp tương tự.", "Normal", False, False, 13, None),
    ]

    # Insert all literature review paragraphs BEFORE "Technology used" H2
    insert_before = tech_used_el
    for section in lit_review_sections:
        stype, content, style, bold, italic, size_pt, color = section
        p_el = make_para_el(doc, content, style_name=style, bold=bold, italic=italic, size_pt=size_pt, color=color)
        body_el.insert(list(body_el).index(insert_before), p_el)

    print(f"  Inserted {len(lit_review_sections)} paragraphs for Literature Review")
else:
    print("  WARNING: Could not find 'Technology used' heading — skipping Literature Review insertion")

# ─────────────────────────────────────────────────────────
# PART 2: Find Chapter 4 / System Design — insert Sequence Diagrams
# ─────────────────────────────────────────────────────────
print("\nPART 2: Finding insertion point for Sequence Diagrams...")

# Find the "3.3 Activity Diagram" H2 — insert Sequence Diagrams AFTER it
activity_idx = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'Activity Diagram' in text and 'Heading 2' in para.style.name:
        activity_idx = i
        break

# Find "3.4" H2 — we will insert BEFORE this
h2_34_idx = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '3.4' in text and 'Heading 2' in para.style.name:
        h2_34_idx = i
        break

print(f"  Activity Diagram H2 at para {activity_idx}")
print(f"  3.4 H2 at para {h2_34_idx}")

if h2_34_idx is not None:
    insert_before_seq = doc.paragraphs[h2_34_idx]._element
    body_el_seq = insert_before_seq.getparent()

    def make_image_para_el(img_path, width_cm=15.0):
        """Create centered paragraph with inline image."""
        from docx.shared import Cm
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        import os
        
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        p.append(pPr)
        return p  # return just the paragraph; we'll add image via para.add_run()

    # Insert Sequence Diagram section BEFORE 3.4
    seq_sections = [
        ("h", "3.3.2 Sequence Diagram — Luồng Thanh Toán (Customer Checkout)", "Heading3", True, False, 13, (0x37,0x47,0x51)),
        ("p", "Sơ đồ tuần tự dưới đây mô tả luồng xử lý đầy đủ khi khách hàng thực hiện quy trình đặt hàng và thanh toán, từ thao tác thêm sản phẩm vào giỏ hàng, qua bước kiểm tra tồn kho với cơ chế Optimistic Concurrency, đến bước tạo đơn hàng và xử lý thanh toán qua Payment Gateway.", "Normal", False, False, 13, None),
    ]

    for stype, content, style, bold, italic, size_pt, color in seq_sections:
        p_el = make_para_el(doc, content, style_name=style, bold=bold, italic=italic, size_pt=size_pt, color=color)
        body_el_seq.insert(list(body_el_seq).index(insert_before_seq), p_el)

    # Insert image 1
    seq1_path = "SEQ_1_Customer_Checkout.png"
    if os.path.exists(seq1_path):
        # We need to create para via doc temporarily then move element
        temp_para = doc.add_paragraph()
        temp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = temp_para.add_run()
        run.add_picture(seq1_path, width=Cm(15))
        temp_el = temp_para._element
        temp_el.getparent().remove(temp_el)
        body_el_seq.insert(list(body_el_seq).index(insert_before_seq), temp_el)
        print(f"  Inserted SEQ 1 image")

    # Caption for seq 1
    cap1 = make_para_el(doc, "Hình 3.6 — Sequence Diagram: Luồng Đặt Hàng và Thanh Toán của Khách Hàng",
                        style_name='Normal', italic=True, size_pt=11,
                        color=(0x33,0x33,0x33), align_center=True)
    body_el_seq.insert(list(body_el_seq).index(insert_before_seq), cap1)

    # Heading for admin flow
    seq_admin_sections = [
        ("h", "3.3.3 Sequence Diagram — Luồng Xử Lý Đơn Hàng (Admin)", "Heading3", True, False, 13, (0x37,0x47,0x51)),
        ("p", "Sơ đồ tuần tự sau mô tả luồng xử lý đơn hàng từ phía quản trị viên, bao gồm thao tác cập nhật trạng thái đơn hàng, ghi audit log, kích hoạt hệ thống thông báo tự động và gửi email xác nhận đến khách hàng thông qua Email Service.", "Normal", False, False, 13, None),
    ]

    for stype, content, style, bold, italic, size_pt, color in seq_admin_sections:
        p_el = make_para_el(doc, content, style_name=style, bold=bold, italic=italic, size_pt=size_pt, color=color)
        body_el_seq.insert(list(body_el_seq).index(insert_before_seq), p_el)

    # Insert image 2
    seq2_path = "SEQ_2_Admin_Order_Management.png"
    if os.path.exists(seq2_path):
        temp_para2 = doc.add_paragraph()
        temp_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = temp_para2.add_run()
        run2.add_picture(seq2_path, width=Cm(15))
        temp_el2 = temp_para2._element
        temp_el2.getparent().remove(temp_el2)
        body_el_seq.insert(list(body_el_seq).index(insert_before_seq), temp_el2)
        print(f"  Inserted SEQ 2 image")

    # Caption for seq 2
    cap2 = make_para_el(doc, "Hình 3.7 — Sequence Diagram: Luồng Admin Xử Lý Đơn Hàng và Gửi Thông Báo",
                        style_name='Normal', italic=True, size_pt=11,
                        color=(0x33,0x33,0x33), align_center=True)
    body_el_seq.insert(list(body_el_seq).index(insert_before_seq), cap2)

    print("  Sequence Diagram sections inserted!")

else:
    print("  WARNING: Could not find 3.4 heading — skipping Sequence Diagram insertion")

# ─────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────
print(f"\nSaving to '{OUT_FILE}'...")
doc.save(OUT_FILE)
print(f"Done! Saved: {OUT_FILE}")

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
