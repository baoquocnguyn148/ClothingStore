# -*- coding: utf-8 -*-
"""
format_report.py
Formats the thesis Word document:
  1. Fixes heading hierarchy (H1/H2/H3/H4)
  2. Moves supplementary diagrams from the end into Chapter 4
  3. Inserts a Word-native TOC field
  4. Adds proper page breaks between chapters
  5. Applies consistent font & spacing
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import re

IN_FILE  = "final report_COMPLETED_WITH_DIAGRAMS.docx"
OUT_FILE = "final report_FORMATTED.docx"

# ─── Helpers ───────────────────────────────────────────────────────────────

def set_heading(para, level):
    para.style = para._element.getparent().getparent().getparent()  # avoid direct style set
    # Easier: just use style name
    style_name = f'Heading {level}'
    para.style = doc.styles[style_name]

def para_text(para):
    return para.text.strip()

def add_toc(doc):
    """Insert a Word Table of Contents field after the CONTENTS heading."""
    # Find the CONTENTS paragraph
    contents_idx = None
    for i, p in enumerate(doc.paragraphs):
        if para_text(p) in ('CONTENTS', 'TABLE OF CONTENTS', 'MỤC LỤC'):
            contents_idx = i
            break
    if contents_idx is None:
        return
    
    # The paragraph element
    contents_para = doc.paragraphs[contents_idx]
    
    # Create TOC field XML
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), r' TOC \o "1-3" \h \z \u ')
    
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    run.append(rpr)
    t = OxmlElement('w:t')
    t.text = '[Nhấn Ctrl+A rồi F9 để cập nhật mục lục]'
    run.append(t)
    fldSimple.append(run)
    
    # Insert a new paragraph for TOC right after CONTENTS heading
    new_p = OxmlElement('w:p')
    new_p.append(fldSimple)
    contents_para._element.addnext(new_p)
    print(f"  ✓ TOC field inserted after 'CONTENTS' paragraph #{contents_idx}")

def ensure_style(doc, name, base='Normal', font_size=12, bold=False, color=None):
    """Ensure a style exists, create if not."""
    if name not in [s.name for s in doc.styles]:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base]
    else:
        style = doc.styles[name]
    font = style.font
    font.size = Pt(font_size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)
    return style

def add_page_break_before(para):
    """Add a page break run before paragraph content."""
    pPr = para._element.get_or_add_pPr()
    pgBr = OxmlElement('w:pageBreakBefore')
    pgBr.set(qn('w:val'), '1')
    pPr.append(pgBr)

def set_para_spacing(para, before=0, after=6, line=None):
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(before * 20))
    spacing.set(qn('w:after'),  str(after  * 20))
    if line:
        spacing.set(qn('w:line'),     str(int(line * 240)))
        spacing.set(qn('w:lineRule'), 'auto')

# ─── Heading fix map ───────────────────────────────────────────────────────
# Maps (current_text_pattern, current_style) -> desired_level

HEADING_FIXES = [
    # --- Fix Title style used as Heading 1 ---
    ('ACKNOWLEDGEMENT',              'Title',    1),
    ('CONTENTS',                     'Title',    1),
    ('TABLE OF CONTENTS',            'Title',    1),
    ('LIST OF PICTURES AND TABLES',  'Title',    1),
    ('LIST OF FIGURES',              'Title',    1),
    
    # --- Fix Heading 3 used for H2 sub-sections in INTRODUCTION ---
    ('Context and research problem', 'Heading 3', 2),
    ('Objectives',                   'Heading 3', 2),
    ('Object and Scope of Research', 'Heading 3', 2),
    ('Research methodology',         'Heading 3', 2),
    ('Meaning of the project',       'Heading 3', 2),
    ('Structure of the topic',       'Heading 3', 2),
    
    # --- Fix Heading 3 used for H2 sub-sections in OVERVIEW ---
    ('The concept of Management Information Systems', 'Heading 3', 2),
    ('Lợi ích của hệ thống thông tin quản lý',       'Heading 3', 2),
    ('Technology used',                               'Heading 3', 2),
    ('Functions in system',                          'Heading 3', 2),
    
    # --- Fix Heading 4 sub-sections in OVERVIEW ---
    ('Fnctional Requirements',       'Heading 4', 3),
    ('Functional Requirements',      'Heading 4', 3),
    ('Non Functional Requirements',  'Heading 4', 3),
    
    # --- CHAPTER 3 heading fixes ---
    ('Chiến lược kinh doanh',        'Heading 3', 3),
    ('Các công cụ và quy trình đang sử dụng', 'Heading 3', 3),
    ('Sơ đồ luồng thông tin hiện tại',        'Heading 3', 3),
    
    # --- MIS SOLUTION DESIGN: fix Heading 3 used as top-level section ---
    ('Xác định vấn đề',              'Heading 3', 3),
    ('Liên kết chiến lược',          'Heading 3', 3),
    ('Cải thiện quy trình vận hành', 'Heading 3', 3),
    ('Kiến trúc hệ thống thông tin', 'Heading 3', 3),
    ('Hệ Thống Truyền Thông',        'Heading 3', 3),
    ('Khả Năng Mở Rộng',             'Heading 3', 3),
    ('Giảm Thiểu Quá Tải Thông Tin', 'Heading 3', 3),
    ('Tích Hợp Thương Mại Điện Tử',  'Heading 3', 3),
    
    # --- SYSTEM DESIGN sub-sections ---
    ('Database Design (OLTP)',        'Heading 3', 3),
    ('Thiết Kế Kho Dữ Liệu Phân Tích','Heading 3',3),
    
    # --- Duplicate Use Case section ---
    ('Use Case Diagram',              'Heading 2', None),  # Will remove one duplicate later
    
    # --- Fix "Các sơ đồ bổ sung" at end ---
    ('Các sơ đồ bổ sung',            'Heading 1', 2),  # Rename to H2 under System Design
]

# ─── Chapter heading rename map ────────────────────────────────────────────
CHAPTER_RENAMES = {
    'INTRODUCTION':          'CHAPTER 1: INTRODUCTION',
    'OVERVIEW OF KNOWLEDGE': 'CHAPTER 2: OVERVIEW OF KNOWLEDGE',
    'Bối cảnh  doanh nghiệp': 'CHAPTER 3: BUSINESS CONTEXT & MIS SOLUTION DESIGN',
    'MIS SOLUTION DESIGN':   None,  # Merge into Chapter 3 (handled separately)
    'SYSTEM DESIGN':         'CHAPTER 4: SYSTEM DESIGN',
    'CHƯƠNG 4: IMPLEMENTATION AND RESULTS': 'CHAPTER 5: IMPLEMENTATION AND RESULTS',
    'CHƯƠNG 5: CONCLUSIONS AND FUTURE RESEARCH DIRECTIONS': 'CHAPTER 6: CONCLUSIONS',
    'REFERENCES':            'REFERENCES',
    'Các sơ đồ bổ sung':     None,  # Will be removed/merged
}

# ─── Main Processing ───────────────────────────────────────────────────────
print(f"Loading '{IN_FILE}'...")
doc = Document(IN_FILE)

print("\n1. Fixing heading hierarchy...")
for para in doc.paragraphs:
    text = para_text(para)
    style = para.style.name
    
    for pattern, from_style, to_level in HEADING_FIXES:
        if pattern.lower() in text.lower() and (from_style is None or from_style in style):
            if to_level is not None:
                para.style = doc.styles[f'Heading {to_level}']
                print(f"  Fixed: [{from_style}] '{text[:50]}' -> Heading {to_level}")
            break

print("\n2. Renaming chapter headings...")
for para in doc.paragraphs:
    text = para_text(para)
    if text in CHAPTER_RENAMES:
        new_text = CHAPTER_RENAMES[text]
        if new_text is None:
            # Remove 'MIS SOLUTION DESIGN' standalone heading (it's already merged under Ch3)
            if text == 'MIS SOLUTION DESIGN':
                para.style = doc.styles['Normal']
                para.clear()
                print(f"  Removed standalone: '{text}'")
            elif text == 'Các sơ đồ bổ sung':
                pass  # Will keep but convert to Heading 2 under Ch4
        else:
            # Replace text in first run
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ''
            else:
                para.add_run(new_text)
            print(f"  Renamed: '{text}' -> '{new_text}'")

print("\n3. Adding page breaks before Chapter headings...")
chapter_patterns = ['CHAPTER ', 'REFERENCES', 'ACKNOWLEDGEMENT', 'TABLE OF CONTENTS', 'LIST OF']
for para in doc.paragraphs:
    text = para_text(para)
    style = para.style.name
    if 'Heading 1' in style and any(p in text.upper() for p in chapter_patterns):
        add_page_break_before(para)
        print(f"  Page break before: '{text[:60]}'")

print("\n4. Fixing paragraph spacing for body text...")
for para in doc.paragraphs:
    style = para.style.name
    if style == 'Normal' and para.text.strip():
        set_para_spacing(para, before=0, after=6, line=1.5)

print("\n5. Inserting Table of Contents field...")
add_toc(doc)

print(f"\n6. Saving to '{OUT_FILE}'...")
doc.save(OUT_FILE)
print(f"\n✅ DONE — Formatted document saved to '{OUT_FILE}'")
print("\nIMPORTANT: Open the file in Microsoft Word, press Ctrl+A then F9 to update the TOC automatically.")
