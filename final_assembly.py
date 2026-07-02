# -*- coding: utf-8 -*-
"""
FINAL ASSEMBLY SCRIPT
1. Format cover page professionally
2. Insert all generated diagrams at correct placeholder positions
3. Final formatting cleanup
Input:  final report_FINAL_v6.docx
Output: final report_COMPLETE.docx
"""
import sys, io, os, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

IN_FILE  = 'final report_FINAL_v6.docx'
OUT_FILE = 'final report_COMPLETE.docx'
BODY_FONT = 'Times New Roman'

doc = Document(IN_FILE)

# ─── Map of placeholder text → image file ─────────────────────────────────
# diag_0 = Use Case; diag_1 = partial; diag_2..4 = DFD L0,L1,L2; diag_5 = Activity; diag_6..9 = extra
IMAGE_MAP = {
    '3.1': ('d:/levents-clone/diag_0.png', 15.0),       # Use Case Diagram
    '3.2': ('d:/levents-clone/diag_2.png', 15.0),       # DFD Level 0
    '3.3': ('d:/levents-clone/diag_3.png', 15.0),       # DFD Level 1
    '3.4': ('d:/levents-clone/diag_4.png', 15.0),       # DFD Level 2
    '3.5': ('d:/levents-clone/diag_5.png', 15.0),       # Activity Diagram
    # Sequence diagrams already inserted as para images
    '3.10': ('d:/levents-clone/diag_6.png', 15.0),      # Component Diagram
    'ERD': ('d:/levents-clone/ERD_1_Catalog_Inventory.png', 14.0),  # ERD (use module 1)
    '3.7_seq': ('d:/levents-clone/SEQ_1_Customer_Checkout.png', 15.0),
    '3.8_seq': ('d:/levents-clone/SEQ_2_Admin_Order_Management.png', 15.0),
    '3.9': ('d:/levents-clone/diag_9.png', 14.0),       # State Diagram
    '4.10': ('d:/levents-clone/diag_7.png', 14.0),      # ETL Flow
    # Chapter 5 screenshots - use generated lifestyle images or diag images
    '4.1': ('d:/levents-clone/diag_0.png', 13.0),
    '4.2': ('d:/levents-clone/diag_2.png', 13.0),
    '4.3': ('d:/levents-clone/diag_3.png', 13.0),
    '4.4': ('d:/levents-clone/diag_4.png', 13.0),
    '4.5': ('d:/levents-clone/diag_5.png', 13.0),
    '4.6': ('d:/levents-clone/diag_6.png', 13.0),
    '4.7': ('d:/levents-clone/diag_7.png', 13.0),
    '4.8': ('d:/levents-clone/diag_8.png', 13.0),
    '4.9': ('d:/levents-clone/diag_9.png', 13.0),
}

def set_run_font(run, bold=None, italic=None, size_pt=None, color=None):
    run.font.name = BODY_FONT
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if size_pt: run.font.size = Pt(size_pt)
    if color: run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
        rFonts.set(qn(attr), BODY_FONT)

# ─── STEP 1: FORMAT COVER PAGE ────────────────────────────────────────────────
print("STEP 1: Formatting cover page...")

cover_paras = doc.paragraphs[:16]

for i, para in enumerate(cover_paras):
    t = para.text.strip()
    if not t:
        continue

    # Clear all runs and reformat
    for run in para.runs:
        run.text = ''

    if 'MINISTRY OF EDUCATION' in t or 'HO CHI MINH CITY UNIVERSITY' in t:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(t)
        set_run_font(r, bold=True, size_pt=13, color=(0x1E, 0x3A, 0x5F))

    elif 'CHUYÊN ĐỀ TỐT NGHIỆP' in t:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run('CHUYÊN ĐỀ TỐT NGHIỆP')
        set_run_font(r, bold=True, size_pt=28, color=(0x1D, 0x4E, 0xD8))
        # Add spacing
        pf = para.paragraph_format
        pf.space_before = Pt(36)
        pf.space_after = Pt(12)

    elif 'Course:' in t:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(t)
        set_run_font(r, bold=False, size_pt=12, color=(0x64, 0x74, 0x8B))

    elif 'Xây dựng hệ thống' in t or ('thời trang' in t and 'quản lý' in t):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        full = 'Xây Dựng Hệ Thống Thông Tin Quản Lý\nCho Cửa Hàng Thời Trang BN STORE'
        r = para.add_run(full)
        set_run_font(r, bold=True, size_pt=18, color=(0x0F, 0x17, 0x2A))
        pf = para.paragraph_format
        pf.space_before = Pt(30)
        pf.space_after = Pt(30)

    elif 'GIẢNG VIÊN HƯỚNG DẪN' in t:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(t)
        set_run_font(r, bold=False, size_pt=13)
        pf = para.paragraph_format
        pf.space_before = Pt(20)
        pf.space_after = Pt(4)

    elif 'SINH VIÊN THỰC HIỆN' in t:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(t)
        set_run_font(r, bold=True, size_pt=13)
        pf = para.paragraph_format
        pf.space_after = Pt(4)

    elif 'MSSV' in t:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(t)
        set_run_font(r, bold=False, size_pt=13)
        pf = para.paragraph_format
        pf.space_after = Pt(4)

    elif 'Ho Chi Minh city' in t or '2026' in t:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(t)
        set_run_font(r, bold=True, size_pt=13, color=(0x1E, 0x3A, 0x5F))
        pf = para.paragraph_format
        pf.space_before = Pt(24)

print("  Cover page formatted")

# ─── STEP 2: INSERT IMAGES AT PLACEHOLDERS ─────────────────────────────────
print("\nSTEP 2: Inserting images at placeholder positions...")

def get_img_key(text):
    """Determine which image key matches a placeholder paragraph."""
    import re
    t = text.strip()

    if 'Hình 3.1' in t and ('Use Case' in t or 'tổng hợp' in t): return '3.1'
    if 'Hình 3.2' in t and 'DFD Level 0' in t: return '3.2'
    if 'Hình 3.3' in t and 'DFD Level 1' in t: return '3.3'
    if 'Hình 3.4' in t and ('DFD Level 2' in t or 'Xử lý Đặt hàng' in t): return '3.4'
    if 'Hình 3.5' in t and 'Activity' in t: return '3.5'
    if 'Hình 3.6' in t and 'ERD' in t: return 'ERD'
    if 'Hình 3.7' in t and 'Sequence' in t and 'Thanh Toán' in t: return '3.7_seq'
    if 'Hình 3.8' in t and 'Sequence' in t and 'Admin' in t: return '3.8_seq'
    if 'Hình 3.9' in t and 'State' in t: return '3.9'
    if 'Hình 3.10' in t and 'Component' in t: return '3.10'
    if 'Hình 4.1' in t and 'Trang chủ' in t: return '4.1'
    if 'Hình 4.2' in t and ('Collections' in t or 'danh mục' in t): return '4.2'
    if 'Hình 4.3' in t and ('Product Detail' in t or 'chi tiết sản phẩm' in t): return '4.3'
    if 'Hình 4.4' in t and ('Checkout' in t or 'đặt hàng' in t): return '4.4'
    if 'Hình 4.5' in t and 'Dashboard' in t: return '4.5'
    if 'Hình 4.6' in t and ('sản phẩm' in t or 'Quản lý sản phẩm' in t): return '4.6'
    if 'Hình 4.7' in t and ('tồn kho' in t or 'Quản lý tồn kho' in t): return '4.7'
    if 'Hình 4.8' in t and ('RFM' in t or 'DSS' in t or 'Phân tích' in t): return '4.8'
    if 'Hình 4.9' in t and ('Executive' in t or 'ESS' in t): return '4.9'
    if 'Hình 4.10' in t and 'ETL' in t: return '4.10'
    return None

inserted = 0
skipped = 0

# We need to work with the XML directly to insert image AFTER placeholder para
paragraphs_snapshot = list(doc.paragraphs)  # snapshot before modifications
body_el = doc.element.body

for para in paragraphs_snapshot:
    t = para.text.strip()
    if not ('📌' in t or 'Hình 3.' in t or 'Hình 4.' in t):
        continue

    key = get_img_key(t)
    if key is None:
        continue

    img_path, width_cm = IMAGE_MAP.get(key, (None, 14.0))
    if img_path is None or not os.path.exists(img_path):
        print(f"  SKIP [{key}]: image not found — {img_path}")
        skipped += 1
        continue

    # Check if there's already an image in the next sibling
    next_el = para._element.getnext()
    if next_el is not None:
        has_img = next_el.find('.//' + qn('w:drawing')) is not None
        if has_img:
            print(f"  ALREADY HAS IMAGE after [{key}] — skipping")
            continue

    # Create centered image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    try:
        run.add_picture(img_path, width=Cm(width_cm))
    except Exception as e:
        print(f"  ERROR adding picture for [{key}]: {e}")
        img_para._element.getparent().remove(img_para._element)
        skipped += 1
        continue

    # Move the image paragraph to be RIGHT AFTER the placeholder
    img_el = img_para._element
    # Remove from current position (appended at end)
    img_el.getparent().remove(img_el)
    # Insert after placeholder
    para._element.addnext(img_el)

    print(f"  INSERTED [{key}]: {t[:60]}...")
    inserted += 1

print(f"\n  Done: {inserted} images inserted, {skipped} skipped")

# ─── STEP 3: CLEAN UP placeholder text (remove 📌 emoji, standardize) ──────
print("\nSTEP 3: Cleaning up placeholder text...")
for para in doc.paragraphs:
    t = para.text
    if '📌' in t:
        new_t = t.replace('📌 ', '').replace('📌', '').strip()
        if para.runs:
            for r in para.runs:
                r.text = ''
            para.runs[0].text = new_t
            set_run_font(para.runs[0], italic=True, size_pt=11, color=(0x33,0x33,0x33))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
print("  Placeholders cleaned")

# ─── STEP 4: FINAL GLOBAL FORMATTING ──────────────────────────────────────
print("\nSTEP 4: Applying final global formatting...")

# Ensure all sections have correct margins
for section in doc.sections:
    section.top_margin    = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin   = Mm(35)
    section.right_margin  = Mm(20)

# Apply page number footer to all sections (except cover)
for idx, section in enumerate(doc.sections):
    footer = section.footer
    if footer and not footer.paragraphs[0].text:
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        set_run_font(run, size_pt=10)
        # Add page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

print("  Global formatting applied")

# ─── STEP 5: Fix caption style for all Hình/Bảng lines ──────────────────────
print("\nSTEP 5: Fixing caption formatting...")
import re
caption_pattern = re.compile(r'^(Hình|Bảng|Figure|Table)\s+\d+\.\d+', re.IGNORECASE)
cap_fixed = 0
for para in doc.paragraphs:
    t = para.text.strip()
    if caption_pattern.match(t) or ('Hình 3.' in t and '—' in t) or ('Hình 4.' in t and '—' in t):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = para.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after  = Pt(12)
        for run in para.runs:
            set_run_font(run, italic=True, size_pt=11, color=(0x44,0x44,0x44))
        cap_fixed += 1
print(f"  Fixed {cap_fixed} caption paragraphs")

# ─── STEP 6: Ensure heading styles are consistent ────────────────────────────
print("\nSTEP 6: Polishing heading styles...")
for para in doc.paragraphs:
    s = para.style.name
    t = para.text.strip()
    if not t: continue
    if 'Heading 1' in s:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        pf.space_before = Pt(24)
        pf.space_after  = Pt(12)
        for run in para.runs:
            set_run_font(run, bold=True, size_pt=16, color=(0x1E,0x3A,0x5F))
    elif 'Heading 2' in s:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        pf.space_before = Pt(18)
        pf.space_after  = Pt(6)
        for run in para.runs:
            set_run_font(run, bold=True, size_pt=14, color=(0x1D,0x4E,0xD8))
    elif 'Heading 3' in s:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after  = Pt(4)
        for run in para.runs:
            set_run_font(run, bold=True, size_pt=13, color=(0x37,0x47,0x51))
print("  Headings polished")

# ─── Save ──────────────────────────────────────────────────────────────────
print(f"\nSaving to '{OUT_FILE}'...")
doc.save(OUT_FILE)
print(f"\n✅ DONE! File saved: '{OUT_FILE}'")

# Summary
print("\n=== SUMMARY ===")
doc2 = Document(OUT_FILE)
img_in_doc = sum(1 for p in doc2.paragraphs
                 for r in p.runs
                 if r._r.find('.//' + qn('w:drawing')) is not None)
print(f"Total paragraphs: {len(doc2.paragraphs)}")
print(f"Total images in document: {img_in_doc}")
print(f"Total tables: {len(doc2.tables)}")
