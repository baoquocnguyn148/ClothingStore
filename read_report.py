# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document

DOCX_FILE = "final report_COMPLETED_WITH_DIAGRAMS.docx"

doc = Document(DOCX_FILE)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print("=" * 60)

for i, para in enumerate(doc.paragraphs):
    style = para.style.name
    text = para.text.strip()
    if text:
        print(f"[{i:4d}] [{style:30s}] {text[:100]}")
