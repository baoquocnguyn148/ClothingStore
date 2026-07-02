# -*- coding: utf-8 -*-
"""
fix_final.py — Final cleanup pass on the formatted doc
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

IN_FILE  = "final report_FORMATTED.docx"
OUT_FILE = "final report_FINAL.docx"

doc = Document(IN_FILE)
paras = doc.paragraphs

print("Pass 1: Fix H4 → H3 for scope sections in Chapter 1...")
for para in paras:
    text = para.text.strip()
    if para.style.name == 'Heading 4' and text in ('Scope of the research', 'Limitations of the topic'):
        para.style = doc.styles['Heading 3']
        print(f"  Fixed H4→H3: '{text}'")

print("\nPass 2: Fix duplicate 'Use Case Diagram' heading...")
uc_count = 0
for para in paras:
    text = para.text.strip()
    if text == 'Use Case Diagram' and para.style.name == 'Heading 2':
        uc_count += 1
        if uc_count == 1:
            # Remove the first occurrence (the one without section number)
            para.style = doc.styles['Normal']
            para.clear()
            print("  Removed duplicate 'Use Case Diagram' heading (first occurrence)")

print("\nPass 3: Fix 'Các sơ đồ bổ sung' — move from REFERENCES section to proper place...")
# Just rename it to something cleaner under System Design
for para in paras:
    text = para.text.strip()
    if text == 'Các sơ đồ bổ sung' and 'Heading' in para.style.name:
        para.style = doc.styles['Heading 2']
        if para.runs:
            para.runs[0].text = '4.6 Additional System Diagrams'
        else:
            para.add_run('4.6 Additional System Diagrams')
        print("  Renamed 'Các sơ đồ bổ sung' → '4.6 Additional System Diagrams'")

print("\nPass 4: Fix section numbers in Chapter 5 (4.x → 5.x)...")
renames_ch5 = {
    '4.1 Môi Trường Triển Khai': '5.1 Môi Trường Triển Khai',
    '4.2 Kết Quả Triển Khai Các Module': '5.2 Kết Quả Triển Khai Các Module',
    '4.2.1 Module 1 — Giao Diện Storefront (B2C)': '5.2.1 Module 1 — Giao Diện Storefront (B2C)',
    '4.2.2 Module 2 — Admin Dashboard (Back Office)': '5.2.2 Module 2 — Admin Dashboard (Back Office)',
    '4.2.3 Module 3 — OLAP Data Warehouse & Power BI': '5.2.3 Module 3 — OLAP Data Warehouse & Power BI',
    '4.3 Kết Quả Kiểm Thử': '5.3 Kết Quả Kiểm Thử',
    '4.3.1 Kiểm thử chức năng (Functional Testing)': '5.3.1 Kiểm thử chức năng (Functional Testing)',
    '4.3.2 Kiểm thử Optimistic Concurrency — Chống Overselling': '5.3.2 Kiểm thử Optimistic Concurrency — Chống Overselling',
    '4.3.3 Kiểm thử Hiệu năng (Performance Testing)': '5.3.3 Kiểm thử Hiệu năng (Performance Testing)',
    '4.3.4 Kiểm thử Bảo mật (Security Testing)': '5.3.4 Kiểm thử Bảo mật (Security Testing)',
    '4.3.5 Kiểm thử Luồng Nghiệp Vụ E2E (End-to-End)': '5.3.5 Kiểm thử Luồng Nghiệp Vụ E2E (End-to-End)',
    '4.4 So Sánh Hiệu Quả Trước và Sau Triển Khai MIS': '5.4 So Sánh Hiệu Quả Trước và Sau Triển Khai MIS',
}
renames_ch6 = {
    '5.1 Kết Luận': '6.1 Kết Luận',
    '5.2 Hạn Chế Của Đề Tài': '6.2 Hạn Chế Của Đề Tài',
    '5.3 Hướng Phát Triển Tiếp Theo': '6.3 Hướng Phát Triển Tiếp Theo',
}
all_renames = {**renames_ch5, **renames_ch6}
for para in paras:
    text = para.text.strip()
    if text in all_renames:
        new_text = all_renames[text]
        if para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ''
        print(f"  Renamed: '{text[:50]}' → '{new_text[:50]}'")

print("\nPass 5: Fix 'Fnctional Requirements' typo...")
for para in paras:
    if 'Fnctional Requirements' in para.text:
        for run in para.runs:
            if 'Fnctional' in run.text:
                run.text = run.text.replace('Fnctional', 'Functional')
                print("  Fixed typo: Fnctional → Functional")

print(f"\nSaving to '{OUT_FILE}'...")
doc.save(OUT_FILE)
print(f"\n✅ DONE — Final document saved: '{OUT_FILE}'")
