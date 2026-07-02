# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import re
import os
import urllib.request
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

MD_FILE = r"C:\Users\usr\.gemini\antigravity\brain\4ccf1c0f-8352-4b02-837a-856255569b69\plantuml_diagrams.md"
DOCX_FILE = "final report_COMPLETED.docx"
OUT_DOCX = "final report_COMPLETED_WITH_DIAGRAMS.docx"

def extract_diagrams(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    blocks = re.findall(r'## HÌNH (3\.\d+) — (.*?)\n.*?```plantuml\n(.*?)\n```', content, re.DOTALL)
    diagrams = []
    for num, title, code in blocks:
        diagrams.append({
            'id': f"Hình {num}",
            'title': title.strip(),
            'code': f"@startuml\n{code.strip()}\n@enduml" if not code.strip().startswith('@startuml') else code.strip()
        })
    return diagrams

def generate_images_kroki(diagrams):
    for i, diag in enumerate(diagrams):
        png_path = f"diag_{i}.png"
        print(f"Generating image for {diag['id']} via Kroki...")
        try:
            req = urllib.request.Request('https://kroki.io/plantuml/png', data=diag['code'].encode('utf-8'), method='POST')
            req.add_header('Content-Type', 'text/plain')
            req.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36')
            with urllib.request.urlopen(req) as response:
                png_data = response.read()
            with open(png_path, 'wb') as f:
                f.write(png_data)
            diag['png'] = png_path
        except Exception as e:
            print(f"Error generating {diag['id']}: {e}")
            diag['png'] = None

def insert_into_docx(diagrams):
    doc = Document(DOCX_FILE)
    marker_map = {diag['id']: diag for diag in diagrams if diag.get('png')}
    
    # 1. Replace existing placeholders
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for marker, diag in list(marker_map.items()):
            if marker in text and 'Hình' in text:
                if i + 1 < len(doc.paragraphs):
                    new_p = doc.paragraphs[i+1].insert_paragraph_before('')
                else:
                    new_p = doc.add_paragraph()
                
                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = new_p.add_run()
                run.add_picture(diag['png'], width=Inches(6.0))
                print(f"Inserted {diag['png']} after marker: {text[:30]}...")
                del marker_map[marker]
                break
                
    # 2. Append remaining diagrams
    if marker_map:
        print("\nAppending missing diagrams to the end of the document...")
        doc.add_page_break()
        doc.add_heading('Các sơ đồ bổ sung', level=1)
        
        # sort map by keys so they append in order
        for marker, diag in sorted(marker_map.items(), key=lambda x: x[0]):
            p = doc.add_paragraph(f"📌 {marker} — {diag['title']}")
            p.style = 'Normal'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            run.add_picture(diag['png'], width=Inches(6.0))
            print(f"Appended {diag['png']} ({marker}) at the end.")

    doc.save(OUT_DOCX)
    print(f"\n[SUCCESS] Saved new document with images to {OUT_DOCX}")

if __name__ == '__main__':
    print("1. Extracting diagrams from markdown...")
    diags = extract_diagrams(MD_FILE)
    print(f"Found {len(diags)} diagrams.")
    
    print("\n2. Generating PNG images via Kroki server...")
    generate_images_kroki(diags)
    
    print("\n3. Inserting into Word document...")
    insert_into_docx(diags)
