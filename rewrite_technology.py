# -*- coding: utf-8 -*-
"""
Rewrite incorrect technology sections in the Word report.
Replace all Oracle/ASP.NET/EF Core references with Next.js/Supabase/PostgreSQL.
Also rewrite Chapter 4 (System Design) content inline.
Input:  final report_FINAL_v5.docx
Output: final report_FINAL_v6.docx
"""
import sys, io, os, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

IN_FILE  = "final report_FINAL_v5.docx"
OUT_FILE = "final report_FINAL_v6.docx"
BODY_FONT = "Times New Roman"

doc = Document(IN_FILE)

# ─── Helper: set fonts on a run ───────────────────────────────────────────────
def set_run_font(run, bold=None, italic=None, size_pt=None, color=None):
    run.font.name = BODY_FONT
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if size_pt: run.font.size = Pt(size_pt)
    if color: run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
        rFonts.set(qn(attr), BODY_FONT)

def replace_para_text(para, new_text, bold=None, italic=None, size_pt=13, color=None):
    """Replace all runs in para with one run containing new_text."""
    for run in para.runs:
        run.text = ''
    if para.runs:
        r = para.runs[0]
        r.text = new_text
        set_run_font(r, bold=bold, italic=italic, size_pt=size_pt, color=color)
    else:
        r = para.add_run(new_text)
        set_run_font(r, bold=bold, italic=italic, size_pt=size_pt, color=color)

# ─── STEP 1: Fix Technology Used section (H2 under Ch2) ───────────────────────
print("STEP 1: Rewriting 'Technology used' section...")

tech_start = None
tech_end = None
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    s = para.style.name
    if 'Technology used' in t and 'Heading 2' in s:
        tech_start = i
    if tech_start and i > tech_start and 'Functions in system' in t and 'Heading 2' in s:
        tech_end = i
        break

print(f"  Tech section: para {tech_start} → {tech_end}")

if tech_start and tech_end:
    # Collect elements to remove (body paragraphs between headings)
    paras_to_remove = list(range(tech_start + 1, tech_end))
    elements_to_remove = [doc.paragraphs[i]._element for i in paras_to_remove]
    body_el = doc.paragraphs[tech_start]._element.getparent()
    insert_before_el = doc.paragraphs[tech_end]._element

    # Remove old paragraphs
    for el in elements_to_remove:
        try:
            body_el.remove(el)
        except:
            pass
    print(f"  Removed {len(elements_to_remove)} old tech paragraphs")

    # Build new technology content
    new_tech_content = [
        # (text, style, bold, italic, size_pt, color)
        ("Hệ thống Thông tin Quản lý BN STORE được xây dựng trên nền tảng công nghệ hiện đại, tận dụng hệ sinh thái JavaScript toàn stack (Full-stack JavaScript) để đảm bảo tốc độ phát triển, hiệu năng cao và khả năng mở rộng linh hoạt. Dưới đây là các công nghệ cốt lõi được sử dụng:", "Normal", False, False, 13, None),

        ("Next.js 15 (React Framework)", "Normal", True, False, 13, (0x1D,0x4E,0xD8)),
        ("Next.js 15 được chọn làm framework chính cho cả frontend và backend (fullstack) của hệ thống. Với kiến trúc App Router và Server Components, Next.js cho phép render trang phía server (SSR) hoặc tĩnh (SSG) tùy theo từng trang — đảm bảo tốc độ tải nhanh và SEO tối ưu cho giao diện storefront B2C. API Routes của Next.js đảm nhận vai trò backend API layer, xử lý các nghiệp vụ như quản lý giỏ hàng, đặt hàng, thanh toán và quản trị hệ thống mà không cần server riêng biệt.", "Normal", False, False, 13, None),

        ("Supabase (Backend-as-a-Service + PostgreSQL)", "Normal", True, False, 13, (0x1D,0x4E,0xD8)),
        ("Supabase là nền tảng Backend-as-a-Service (BaaS) mã nguồn mở đóng vai trò hạ tầng dữ liệu trung tâm của toàn bộ hệ thống. Supabase cung cấp một bộ dịch vụ tích hợp hoàn chỉnh bao gồm: (1) Cơ sở dữ liệu PostgreSQL với 38 bảng nghiệp vụ được thiết kế đầy đủ ràng buộc khóa ngoại và index tối ưu; (2) Supabase Auth cho xác thực người dùng đa phương thức (email/password, OAuth); (3) Supabase Storage cho lưu trữ ảnh sản phẩm và tài nguyên; (4) Supabase Realtime cho cập nhật dữ liệu real-time trên Admin Dashboard thông qua WebSocket; (5) Row Level Security (RLS) để kiểm soát quyền truy cập dữ liệu cấp hàng (row-level) mà không cần viết middleware xác thực phức tạp.", "Normal", False, False, 13, None),

        ("PostgreSQL (Cơ sở dữ liệu quan hệ)", "Normal", True, False, 13, (0x1D,0x4E,0xD8)),
        ("PostgreSQL 15 là hệ quản trị cơ sở dữ liệu quan hệ (RDBMS) mã nguồn mở mạnh mẽ, được Supabase sử dụng làm database engine. PostgreSQL cung cấp các tính năng nâng cao như: hỗ trợ kiểu dữ liệu JSONB cho dữ liệu bán cấu trúc (địa chỉ giao hàng, cấu hình khuyến mãi); UUID làm khóa chính; Full-Text Search (FTS) với tsvector cho tìm kiếm sản phẩm tiếng Việt; và các trigger tự động cập nhật updated_at. Toàn bộ 38 bảng của BN STORE được triển khai trên PostgreSQL với các constraint, index và foreign key được định nghĩa chặt chẽ.", "Normal", False, False, 13, None),

        ("TypeScript", "Normal", True, False, 13, (0x1D,0x4E,0xD8)),
        ("TypeScript được sử dụng xuyên suốt toàn bộ codebase (frontend và backend API routes) nhằm đảm bảo type safety, giảm thiểu lỗi runtime và nâng cao khả năng bảo trì. Các kiểu dữ liệu (type definitions) được generate tự động từ Supabase Schema thông qua Supabase CLI, đảm bảo sự nhất quán giữa database schema và application code.", "Normal", False, False, 13, None),

        ("Tailwind CSS", "Normal", True, False, 13, (0x1D,0x4E,0xD8)),
        ("Tailwind CSS là framework CSS utility-first được sử dụng để xây dựng toàn bộ giao diện người dùng. Với cách tiếp cận atomic CSS, Tailwind cho phép xây dựng giao diện responsive, thẩm mỹ cao và nhất quán mà không cần viết CSS tùy chỉnh. Giao diện storefront B2C được thiết kế theo phong cách tối giản (minimalist) hiện đại, tối ưu cho trải nghiệm mua sắm trên cả desktop lẫn mobile.", "Normal", False, False, 13, None),

        ("Vercel (Deployment Platform)", "Normal", True, False, 13, (0x1D,0x4E,0xD8)),
        ("Vercel là nền tảng triển khai (deployment platform) được tối ưu hóa đặc biệt cho ứng dụng Next.js. Hệ thống BN STORE được triển khai lên Vercel với Edge Network phân tán toàn cầu, đảm bảo thời gian phản hồi thấp. Vercel tự động xử lý CI/CD pipeline, preview deployments cho mỗi pull request và scaling tự động theo tải.", "Normal", False, False, 13, None),

        ("Resend (Email Service)", "Normal", True, False, 13, (0x1D,0x4E,0xD8)),
        ("Resend là dịch vụ gửi email giao dịch (transactional email) được tích hợp để gửi email xác nhận đơn hàng, thông báo trạng thái vận chuyển và email marketing. API của Resend được gọi từ Next.js API Routes với các template email được thiết kế chuẩn HTML responsive.", "Normal", False, False, 13, None),
    ]

    for text, style, bold, italic, size_pt, color in new_tech_content:
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style.replace(' ', ''))
        pPr.append(pStyle)
        p.append(pPr)

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
            rFonts.set(qn(attr), BODY_FONT)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(sz)
        if bold:
            rPr.append(OxmlElement('w:b'))
        if italic:
            rPr.append(OxmlElement('w:i'))
        if color:
            clr = OxmlElement('w:color')
            clr.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color))
            rPr.append(clr)
        r.append(rPr)
        t_el = OxmlElement('w:t')
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_el.text = text
        r.append(t_el)
        p.append(r)

        body_el.insert(list(body_el).index(insert_before_el), p)

    print(f"  Inserted {len(new_tech_content)} new technology paragraphs")

# ─── STEP 2: Fix incorrect tech references in body text ────────────────────────
print("\nSTEP 2: Replacing incorrect tech references in all body paragraphs...")

replacements = [
    # Wrong → Correct
    ("Oracle Database 11g", "PostgreSQL (Supabase)"),
    ("Oracle Database", "PostgreSQL (Supabase)"),
    ("Oracle", "Supabase PostgreSQL"),
    ("ASP.NET Core Identity", "Supabase Auth"),
    ("ASP.NET Core", "Next.js"),
    ("ASP.NET", "Next.js"),
    ("Entity Framework Core", "Supabase JS Client"),
    ("Entity Framework", "Supabase JS Client"),
    ("analytics.db", "Supabase Analytics Schema"),
    ("SQLite", "Supabase (PostgreSQL)"),
    (".NET", "Next.js"),
    ("C#", "TypeScript"),
    ("POCO", "TypeScript Interface"),
    ("Windows Task Scheduler", "Vercel Cron Jobs"),
    ("Code First Migration", "Supabase Migrations"),
    ("IPN Webhook", "Supabase Webhook / Next.js API Route"),
    ("Bootstrap 5", "Tailwind CSS"),
    ("jQuery AJAX", "React Server Actions / fetch()"),
    ("ASP.NET MVC", "Next.js App Router"),
]

fixed_count = 0
for para in doc.paragraphs:
    original = para.text
    needs_fix = any(old in original for old, new in replacements)
    if needs_fix:
        new_text = original
        for old, new in replacements:
            new_text = new_text.replace(old, new)
        if new_text != original:
            replace_para_text(para, new_text)
            fixed_count += 1

print(f"  Fixed {fixed_count} paragraphs with incorrect technology references")

# ─── STEP 3: Rewrite Chapter 4 Data Layer section ─────────────────────────────
print("\nSTEP 3: Rewriting Chapter 4 - Data Layer section...")

# Find and replace the Data Layer / OLAP section headings and content
for para in doc.paragraphs:
    t = para.text.strip()
    s = para.style.name

    # Fix "Database Design (OLTP)" heading content
    if 'Database Design (OLTP)' in t and 'Heading' in s:
        # Content is already below it – just ensure heading is correct
        pass

    # Fix the OLAP/Warehouse section
    if 'Thiết Kế Kho Dữ Liệu Phân Tích' in t and 'Heading' in s:
        # Already good heading - content below will be fixed by step 2 replacements
        pass

# Specifically fix the OLAP description paragraphs
olap_paras_fixed = 0
for para in doc.paragraphs:
    t = para.text
    if 'Oracle' in t or 'SQLite' in t or 'analytics.db' in t or 'Windows Task Scheduler' in t:
        new_t = t
        new_t = new_t.replace('Oracle', 'Supabase PostgreSQL (OLTP)')
        new_t = new_t.replace('SQLite', 'Supabase Analytics Schema')
        new_t = new_t.replace('analytics.db', 'analytics schema trong Supabase')
        new_t = new_t.replace('Windows Task Scheduler', 'Vercel Cron Jobs (scheduled at 2:00 AM UTC+7)')
        if new_t != t:
            replace_para_text(para, new_t)
            olap_paras_fixed += 1

print(f"  Fixed {olap_paras_fixed} OLAP/Data Layer paragraphs")

# ─── STEP 4: Fix Chapter 3 architecture description ──────────────────────────
print("\nSTEP 4: Fixing Chapter 3 architecture description...")

arch_fixed = 0
for para in doc.paragraphs:
    t = para.text
    if 'Clean Architecture' in t and ('ASP' in t or '.NET' in t or 'Oracle' in t):
        new_t = t.replace('ASP.NET Core Identity', 'Supabase Auth')
        new_t = new_t.replace('ASP.NET Core', 'Next.js')
        new_t = new_t.replace('Oracle', 'Supabase PostgreSQL')
        new_t = new_t.replace('Entity Framework Core Code  First Migration', 'Supabase Migrations CLI')
        if new_t != t:
            replace_para_text(para, new_t)
            arch_fixed += 1

# Fix the "Kiến trúc hệ thống thông tin" description
for para in doc.paragraphs:
    t = para.text
    if 'Oracle Database 11g với 16 bảng' in t:
        new_t = t.replace(
            'Cơ sở dữ liệu vận hành (OLTP) sử dụng Oracle Database 11g với 16 bảng nghiệp vụ được tối ưu hóa cho các g',
            'Cơ sở dữ liệu vận hành (OLTP) sử dụng PostgreSQL qua Supabase với 38 bảng nghiệp vụ được tối ưu hóa cho các giao dịch thời gian thực.'
        )
        replace_para_text(para, new_t)
        arch_fixed += 1

print(f"  Fixed {arch_fixed} architecture description paragraphs")

# ─── STEP 5: Rewrite Data Layer section (inside Chapter 4) to match Supabase ──
print("\nSTEP 5: Rewriting Chapter 4 Data Layer intro...")

for para in doc.paragraphs:
    t = para.text.strip()
    s = para.style.name
    # Find the paragraph describing the OLTP layer mentioning Oracle
    if ('vận hành (OLTP) sử dụng' in t or 'Oracle' in t) and 'Heading' not in s:
        new_t = re.sub(
            r'sử dụng (Oracle Database \d+g|Oracle)',
            'sử dụng PostgreSQL 15 thông qua Supabase',
            t
        )
        new_t = new_t.replace('16 bảng', '38 bảng')
        if new_t != t:
            replace_para_text(para, new_t)

# Fix Data Dictionary references
for para in doc.paragraphs:
    t = para.text
    if 'Data Dictionary' in t and 'Bảng Inventories' in t:
        new_t = t.replace('Bảng Inventories (★ = ', 'Bảng product_variants (★ = ')
        if new_t != t:
            replace_para_text(para, new_t)
    if 'Data Dictionary: Bảng Products' in t:
        new_t = t.replace('Bảng Products', 'Bảng products (PostgreSQL/Supabase)')
        if new_t != t:
            replace_para_text(para, new_t)
    if 'Data Dictionary: Bảng ProductSkus' in t:
        new_t = t.replace('Bảng ProductSkus', 'Bảng product_variants (SKUs)')
        if new_t != t:
            replace_para_text(para, new_t)
    if 'Data Dictionary: Bảng Orders' in t:
        new_t = t.replace('Bảng Orders', 'Bảng orders (PostgreSQL/Supabase)')
        if new_t != t:
            replace_para_text(para, new_t)

# ─── STEP 6: Fix Optimistic Concurrency section (EF Core → Supabase/PostgreSQL) ──
print("\nSTEP 6: Fixing Optimistic Concurrency section...")

for para in doc.paragraphs:
    t = para.text
    if 'Entity Framework Core' in t and 'Optimistic' in t:
        new_t = t.replace(
            'Entity Framework Core cung cấp cơ chế Optimistic Concurrency thông qua thuộc tính [Timestamp] (hay ConcurrencyToken)',
            'Supabase (PostgreSQL) hỗ trợ cơ chế Optimistic Concurrency thông qua cột version (integer) hoặc updated_at (timestamp)'
        )
        new_t = new_t.replace('RowVersion', 'version column')
        new_t = new_t.replace('[Timestamp]', 'version')
        new_t = new_t.replace('ConcurrencyToken', 'optimistic lock column')
        replace_para_text(para, new_t)

    if 'RowVersion' in t and 'Heading' not in para.style.name:
        new_t = t.replace('RowVersion', 'version')
        if new_t != t:
            replace_para_text(para, new_t)

# Fix section heading about Optimistic Concurrency
for para in doc.paragraphs:
    if '3.4.2 Giải Pháp: Optimistic Concurrency với RowVersion' in para.text:
        replace_para_text(para, '3.4.2 Giải Pháp: Optimistic Concurrency với Version Column trong Supabase PostgreSQL',
                          bold=True, size_pt=13, color=(0x37,0x47,0x51))

print("  Done fixing concurrency section")

# ─── STEP 7: Fix Implementation Chapter (Chapter 5) ─────────────────────────
print("\nSTEP 7: Fixing Chapter 5 (Implementation) ...")
impl_fixed = 0
for para in doc.paragraphs:
    t = para.text
    changes = [
        ("Module 3 — OLAP Data Warehouse & Power BI", "Module 3 — Analytics Dashboard & Báo Cáo Kinh Doanh"),
        ("Microsoft Power BI", "Recharts / Biểu đồ tích hợp Next.js"),
        ("Power BI", "Recharts Analytics Dashboard"),
        ("analytics.db", "Supabase Analytics Schema"),
        ("SQLite analytics", "Supabase Analytics"),
        ("ETL tự động mỗi đêm lúc 2:00 AM theo Windows Task Scheduler", "ETL tự động theo Vercel Cron Jobs (2:00 AM UTC+7)"),
        ("Oracle sang kho dữ liệu phân tích SQLite", "OLTP sang Analytics Schema trong Supabase"),
        ("Oracle sang", "Supabase OLTP sang"),
        ("Chuẩn hóa ngày tháng về múi giờ UTC+7", "Chuẩn hóa timestamp về múi giờ UTC+7 (Asia/Ho_Chi_Minh)"),
        ("INSERT OR REPLACE", "UPSERT (ON CONFLICT DO UPDATE)"),
    ]
    new_t = t
    for old, new in changes:
        new_t = new_t.replace(old, new)
    if new_t != t:
        replace_para_text(para, new_t)
        impl_fixed += 1

print(f"  Fixed {impl_fixed} implementation paragraphs")

# ─── STEP 8: Fix Module descriptions ────────────────────────────────────────
print("\nSTEP 8: Fixing module/screenshot descriptions...")
for para in doc.paragraphs:
    t = para.text
    # Fix Module 3 description
    if 'OLAP Data Warehouse' in t and ('5.2.3' in t or 'Module 3' in t) and 'Heading' in para.style.name:
        replace_para_text(para,
            t.replace('OLAP Data Warehouse & Power BI', 'Analytics Dashboard & Báo Cáo'),
            bold=True, size_pt=13, color=(0x37,0x47,0x51))

# ─── Save ────────────────────────────────────────────────────────────────────
print(f"\nSaving to '{OUT_FILE}'...")
doc.save(OUT_FILE)
print(f"\n✅ DONE — Saved: '{OUT_FILE}'")
print("\nSummary of changes:")
print("  ✓ Rewrote 'Technology used' section — Oracle/ASP.NET → Next.js/Supabase/PostgreSQL")
print("  ✓ Fixed all body text references to Oracle, Entity Framework, SQLite, ASP.NET")
print("  ✓ Fixed Optimistic Concurrency section — RowVersion → version column (PostgreSQL)")
print("  ✓ Fixed Chapter 5 Implementation — Power BI → Recharts, Windows Scheduler → Vercel Cron")
print("  ✓ Fixed Data Dictionary table descriptions")
